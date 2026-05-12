"""
Template-driven DOCX builder.

Pipeline:
  1. TemplateAnalyzer.analyze(path) -> DocumentStructure
     Reads the template .docx, captures every paragraph's text + styles,
     and detects "sample" variable values (client name, fee, NIE).

  2. FieldSubstitutor.fill(structure, fields) -> DocumentStructure
     Replaces sample variable values with actual extracted field values.

  3. DocxRenderer.render(structure, output_path)
     Builds a new .docx that mirrors the template formatting
     with all variable slots filled in.
"""
from __future__ import annotations

import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from lxml import etree


# ── Spanish month names ───────────────────────────────────────────────────────

_MONTHS_ES = [
    "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
]


def _today_es() -> str:
    d = datetime.now()
    return f"En Madrid, a {d.day} de {_MONTHS_ES[d.month - 1]} de {d.year}."


# ── Regex helpers ─────────────────────────────────────────────────────────────

# Matches "MIL DOSCIENTOS EUROS (1.200,00 €)" or plain "1.200,00 €"
_MONEY_RE = re.compile(
    r"[A-ZÁÉÍÓÚÜÑ][A-ZÁÉÍÓÚÜÑ\s]{3,}\s*\(\s*[\d\.]+,\d{2}\s*€\)"
    r"|[\d\.]+,\d{2}\s*€",
    re.UNICODE,
)
_NIE_RE = re.compile(r"\b[XYZ]\d{7}[A-Z]\b")
_CAPS_NAME_RE = re.compile(r"[A-ZÁÉÍÓÚÜÑ]{2,}(?:\s+[A-ZÁÉÍÓÚÜÑ]{2,}){1,4}")

_CJK_RE = re.compile(r"[一-鿿㐀-䶿぀-ヿ]")


def _has_cjk(text: str) -> bool:
    return bool(_CJK_RE.search(text))


_CLAUSE_HEADING_RE = re.compile(
    r"^(PRIMERA|SEGUNDA|TERCERA|CUARTA|QUINTA|SEXTA|S[EÉ]PTIMA|OCTAVA|NOVENA"
    r"|D[EÉ]CIMA|UNDÉCIMA|DUODÉCIMA|DECIMOTERCERA)",
    re.IGNORECASE,
)


# ── Data model ────────────────────────────────────────────────────────────────

@dataclass
class RunStyle:
    text: str
    bold: bool = False
    italic: bool = False
    underline: bool = False
    font_name: str = "Times New Roman"
    font_size_pt: float = 11.0
    color_hex: str | None = None
    all_caps: bool = False


@dataclass
class BlockStyle:
    alignment: str = "LEFT"       # LEFT | CENTER | RIGHT | JUSTIFY
    space_before_pt: float = 0.0
    space_after_pt: float = 0.0
    indent_left_cm: float = 0.0
    hanging_cm: float = 0.0
    border_bottom: bool = False
    border_color: str = "1F3864"


@dataclass
class Block:
    kind: str                     # title | subtitle | section_heading | clause_title |
                                  # date_line | body | list_item | divider | empty | signature
    runs: list[RunStyle]
    style: BlockStyle
    raw_text: str                 # full concatenated paragraph text


@dataclass
class DocumentStructure:
    blocks: list[Block]
    sample_client_name: str = ""  # ALL-CAPS name found in template
    sample_fee: str = ""          # money string found in template (e.g. "1.200,00 €")
    sample_nie: str = ""          # NIE found in template


# ── TemplateAnalyzer ──────────────────────────────────────────────────────────

class TemplateAnalyzer:
    """Parse a .docx template → DocumentStructure (blocks with captured styles)."""

    def analyze(self, path: Path) -> DocumentStructure:
        doc = Document(path)
        blocks: list[Block] = []

        for para in doc.paragraphs:
            blocks.append(self._para_to_block(para))

        # Detect signature table
        for table in doc.tables:
            combined = " ".join(c.text for r in table.rows for c in r.cells).upper()
            if any(kw in combined for kw in ("DESPACHO", "FIRMA", "_____")):
                blocks.append(Block(
                    kind="signature",
                    runs=[],
                    style=BlockStyle(),
                    raw_text="[SIGNATURE]",
                ))
                break

        structure = DocumentStructure(blocks=blocks)
        self._detect_sample_values(structure)
        return structure

    # ── paragraph → Block ──

    def _para_to_block(self, para) -> Block:
        raw = para.text.strip()
        style = self._capture_style(para)
        runs = self._capture_runs(para)
        kind = self._classify(raw, para, runs)
        return Block(kind=kind, runs=runs, style=style, raw_text=raw)

    def _capture_style(self, para) -> BlockStyle:
        pf = para.paragraph_format

        align_map = {
            WD_ALIGN_PARAGRAPH.CENTER: "CENTER",
            WD_ALIGN_PARAGRAPH.RIGHT: "RIGHT",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "JUSTIFY",
        }
        alignment = align_map.get(para.alignment, "LEFT")

        def to_pt(length) -> float:
            try:
                return float(length.pt) if length else 0.0
            except Exception:
                return 0.0

        def to_cm(length) -> float:
            try:
                return float(length.cm) if length else 0.0
            except Exception:
                return 0.0

        hanging = 0.0
        if pf.first_line_indent and pf.first_line_indent < 0:
            hanging = abs(to_cm(pf.first_line_indent))

        border_bottom = False
        border_color = "1F3864"
        pPr = para._p.pPr
        if pPr is not None:
            pBdr = pPr.find(qn("w:pBdr"))
            if pBdr is not None:
                bottom = pBdr.find(qn("w:bottom"))
                if bottom is not None:
                    border_bottom = True
                    border_color = bottom.get(qn("w:color"), "1F3864")

        return BlockStyle(
            alignment=alignment,
            space_before_pt=to_pt(pf.space_before),
            space_after_pt=to_pt(pf.space_after),
            indent_left_cm=to_cm(pf.left_indent),
            hanging_cm=hanging,
            border_bottom=border_bottom,
            border_color=border_color,
        )

    def _capture_runs(self, para) -> list[RunStyle]:
        result: list[RunStyle] = []
        for run in para.runs:
            if not run.text:
                continue
            font = run.font
            size = 11.0
            try:
                if font.size:
                    size = float(font.size.pt)
            except Exception:
                pass
            color = None
            try:
                if font.color and font.color.type is not None and font.color.rgb:
                    color = str(font.color.rgb)
            except Exception:
                pass
            result.append(RunStyle(
                text=run.text,
                bold=bool(run.bold),
                italic=bool(run.italic),
                underline=bool(run.underline),
                font_name=font.name or "Times New Roman",
                font_size_pt=size,
                color_hex=color,
                all_caps=bool(font.all_caps),
            ))
        if not result and para.text:
            result.append(RunStyle(text=para.text))
        return result

    def _classify(self, text: str, para, runs: list[RunStyle]) -> str:
        if not text:
            pPr = para._p.pPr
            if pPr is not None and pPr.find(qn("w:pBdr")) is not None:
                return "divider"
            return "empty"

        upper = text.strip().upper()
        is_centered = para.alignment == WD_ALIGN_PARAGRAPH.CENTER
        is_bold = any(r.bold for r in runs)
        max_size = max((r.font_size_pt for r in runs), default=11.0)

        if is_centered and is_bold:
            if max_size >= 16:
                return "title"
            if max_size >= 13:
                return "subtitle"

        if upper in {"REUNIDOS", "EXPONEN", "CLÁUSULAS", "CLAUSULAS"}:
            return "section_heading"

        if _CLAUSE_HEADING_RE.match(upper):
            return "clause_title"

        low = text.lower()
        if low.startswith("en ") and (
            "de " in low
            or para.alignment == WD_ALIGN_PARAGRAPH.RIGHT
            or is_centered
        ):
            return "date_line"

        if re.match(r"^([a-z]\)|—|–|-|—)", text.strip()):
            return "list_item"

        return "body"

    # ── detect sample variable values in the template ──

    def _detect_sample_values(self, structure: DocumentStructure) -> None:
        for block in structure.blocks:
            raw = block.raw_text

            # NIE
            if not structure.sample_nie:
                m = _NIE_RE.search(raw)
                if m:
                    structure.sample_nie = m.group()

            # Fee: look for "1.200,00 €" pattern
            if not structure.sample_fee:
                m = re.search(r"[\d\.]+,\d{2}\s*€", raw)
                if m:
                    structure.sample_fee = m.group().strip()

            # Client name: bold ALL-CAPS sequence in client-related paragraphs.
            # Check both "de una parte" (client) and "de otra parte" (despacho) blocks,
            # plus any subtitle, to maximise detection coverage.
            if not structure.sample_client_name and block.kind in ("body", "subtitle"):
                low = raw.lower()
                in_client_ctx = (
                    "de una parte" in low
                    or "de otra parte" in low
                    or "el cliente" in low
                    or "don/doña" in low
                    or "don/" in low
                    or block.kind == "subtitle"
                )
                if in_client_ctx:
                    for run in block.runs:
                        if run.bold:
                            m = _CAPS_NAME_RE.search(run.text)
                            if m and len(m.group().split()) >= 2:
                                structure.sample_client_name = m.group()
                                break


# ── FieldSubstitutor ──────────────────────────────────────────────────────────

class FieldSubstitutor:
    """Replace sample variable values in DocumentStructure with actual field values."""

    def fill(self, structure: DocumentStructure, fields: dict) -> DocumentStructure:
        new_name = self._get(
            fields,
            # AI 常见输出 key 全覆盖
            "nombre_cliente", "nombre", "name", "cliente", "client",
            "nombre_completo", "full_name", "apellidos", "client_name",
            "nombre_del_cliente", "customer_name",
        )
        fee_raw = self._get(
            fields,
            "honorarios", "price", "precio", "fee",
            "importe", "amount", "tarifa", "honorario",
        )
        new_fee = self._format_fee(fee_raw)
        new_nie = self._get(fields, "nie", "nif", "dni")

        filled: list[Block] = [
            self._fill_block(
                block,
                sample_name=structure.sample_client_name,
                new_name=new_name,
                sample_fee=structure.sample_fee,
                new_fee=new_fee,
                sample_nie=structure.sample_nie,
                new_nie=new_nie,
            )
            for block in structure.blocks
        ]

        return DocumentStructure(
            blocks=filled,
            sample_client_name=new_name or structure.sample_client_name,
            sample_fee=new_fee or structure.sample_fee,
            sample_nie=new_nie or structure.sample_nie,
        )

    # ── helpers ──

    def _get(self, fields: dict, *keys: str, default: str = "") -> str:
        for k in keys:
            v = fields.get(k)
            if v and str(v).strip().lower() not in {"none", "null", "n/a", ""}:
                return str(v).strip()
        return default

    def _format_fee(self, raw: str) -> str:
        """Normalise any fee string to Spanish format: '1.200,00 €'."""
        if not raw:
            return ""
        # Strip units and whitespace
        cleaned = re.sub(r"(?i)(euros?|eur|€|\s)", "", raw.strip())
        # Detect decimal separator
        if re.search(r",\d{2}$", cleaned):
            # European "1.200,00" → remove thousand dots, swap comma
            cleaned = cleaned.replace(".", "").replace(",", ".")
        elif re.search(r"\.\d{2}$", cleaned):
            # US/EN "1,200.00" → remove commas
            cleaned = cleaned.replace(",", "")
        else:
            # Plain integer "1200"
            cleaned = re.sub(r"[,\.]", "", cleaned)
        try:
            amount = float(cleaned)
            int_part = f"{int(amount):,}".replace(",", ".")
            frac_str = f"{amount % 1:.2f}"[2:]
            return f"{int_part},{frac_str} €"
        except (ValueError, TypeError):
            return raw

    # ── block-level fill ──

    def _fill_block(
        self,
        block: Block,
        *,
        sample_name: str,
        new_name: str,
        sample_fee: str,
        new_fee: str,
        sample_nie: str,
        new_nie: str,
    ) -> Block:
        # Passthrough for structural elements
        if block.kind in ("signature", "divider", "empty"):
            return block

        # Date line: always overwrite with today
        if block.kind == "date_line":
            proto = block.runs[0] if block.runs else RunStyle(text="")
            new_run = RunStyle(
                text=_today_es(),
                bold=proto.bold,
                italic=proto.italic,
                font_name=proto.font_name,
                font_size_pt=proto.font_size_pt,
                color_hex=proto.color_hex,
            )
            return Block(
                kind=block.kind,
                runs=[new_run],
                style=block.style,
                raw_text=_today_es(),
            )

        # Run-level substitution
        new_runs = [
            self._fill_run(
                run,
                sample_name=sample_name, new_name=new_name,
                sample_fee=sample_fee, new_fee=new_fee,
                sample_nie=sample_nie, new_nie=new_nie,
            )
            for run in block.runs
        ]
        return Block(
            kind=block.kind,
            runs=new_runs,
            style=block.style,
            raw_text="".join(r.text for r in new_runs),
        )

    def _fill_run(
        self,
        run: RunStyle,
        *,
        sample_name: str,
        new_name: str,
        sample_fee: str,
        new_fee: str,
        sample_nie: str,
        new_nie: str,
    ) -> RunStyle:
        text = run.text

        # 1. Client name
        if new_name:
            if sample_name and sample_name in text:
                # Direct replacement of known sample name
                text = text.replace(sample_name, new_name.upper())
            elif not sample_name and run.bold:
                # Fallback A: bold run has ALL-CAPS Latin name → replace it
                if _CAPS_NAME_RE.search(text):
                    text = _CAPS_NAME_RE.sub(new_name.upper(), text, count=1)
                # Fallback B: bold run already contains a CJK name (previous client) → replace it
                elif _has_cjk(text) and _CJK_RE.search(text):
                    text = _CJK_RE.sub("", text).strip()  # remove old CJK, then append new
                    text = (text + " " + new_name).strip() if text else new_name

        # 2. NIE
        if new_nie:
            if sample_nie and sample_nie in text:
                text = text.replace(sample_nie, new_nie.upper())
            elif _NIE_RE.search(text):
                text = _NIE_RE.sub(new_nie.upper(), text)

        # 3. Fee — long form first ("MIL DOSCIENTOS EUROS (…)"), then plain amount
        if new_fee:
            text = re.sub(
                r"[A-ZÁÉÍÓÚÜÑ][A-ZÁÉÍÓÚÜÑ\s]{3,}\s*\(\s*[\d\.]+,\d{2}\s*€\)",
                new_fee,
                text,
                flags=re.UNICODE,
            )
            text = re.sub(r"[\d\.]+,\d{2}\s*€", new_fee, text)

        return RunStyle(
            text=text,
            bold=run.bold,
            italic=run.italic,
            underline=run.underline,
            font_name=run.font_name,
            font_size_pt=run.font_size_pt,
            color_hex=run.color_hex,
            all_caps=run.all_caps,
        )


# ── DocxRenderer ──────────────────────────────────────────────────────────────

class DocxRenderer:
    """Build a new .docx from a DocumentStructure, preserving all captured styles."""

    def render(self, structure: DocumentStructure, output_path: Path) -> None:
        doc = Document()
        # python-docx adds one empty paragraph by default — remove it
        for p in list(doc.paragraphs):
            self._remove(p._element)

        for block in structure.blocks:
            if block.kind == "signature":
                self._add_signature(doc, structure.sample_client_name)
            elif block.kind == "divider":
                self._add_divider(doc, block.style)
            elif block.kind == "empty":
                self._add_empty(doc)
            else:
                self._add_block(doc, block)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)

    # ── renderers ──

    def _add_block(self, doc: Document, block: Block) -> None:
        para = doc.add_paragraph()
        self._apply_para_style(para, block.style)
        for rs in block.runs:
            if rs.text:
                self._apply_run_style(para.add_run(rs.text), rs)

    def _add_divider(self, doc: Document, style: BlockStyle) -> None:
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(style.space_before_pt or 8)
        para.paragraph_format.space_after = Pt(style.space_after_pt or 8)
        self._set_bottom_border(para, style.border_color or "1F3864")

    def _add_empty(self, doc: Document) -> None:
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)

    def _add_signature(self, doc: Document, client_name: str) -> None:
        table = doc.add_table(rows=4, cols=2)
        # Strip all borders from table
        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = etree.SubElement(tbl, qn("w:tblPr"))
        tblBorders = etree.SubElement(tblPr, qn("w:tblBorders"))
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = etree.SubElement(tblBorders, qn(f"w:{edge}"))
            b.set(qn("w:val"), "none")

        rows_data = [
            ("EL DESPACHO", "EL CLIENTE", True, False),
            ("", "", False, False),
            ("______________________________", "______________________________", False, False),
            ("Firma y sello del Despacho", client_name or "EL CLIENTE", False, True),
        ]
        for row_idx, (left, right, header_bold, name_bold) in enumerate(rows_data):
            row = table.rows[row_idx]
            for col_idx, label in enumerate((left, right)):
                cell = row.cells[col_idx]
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(label)
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)
                run.bold = header_bold or (name_bold and col_idx == 1)
                if _has_cjk(label):
                    self._set_east_asian_font(run, "Microsoft YaHei")

    # ── style helpers ──

    def _apply_para_style(self, para, style: BlockStyle) -> None:
        pf = para.paragraph_format
        if style.space_before_pt:
            pf.space_before = Pt(style.space_before_pt)
        if style.space_after_pt:
            pf.space_after = Pt(style.space_after_pt)
        if style.indent_left_cm:
            pf.left_indent = Cm(style.indent_left_cm)
        if style.hanging_cm:
            pf.first_line_indent = Cm(-style.hanging_cm)

        para.alignment = {
            "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
            "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
            "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }.get(style.alignment, WD_ALIGN_PARAGRAPH.LEFT)

        if style.border_bottom:
            self._set_bottom_border(para, style.border_color)

    def _apply_run_style(self, run, rs: RunStyle) -> None:
        run.bold = rs.bold
        run.italic = rs.italic
        run.underline = rs.underline
        run.font.name = rs.font_name
        if rs.font_size_pt:
            run.font.size = Pt(rs.font_size_pt)
        if rs.all_caps:
            run.font.all_caps = True
        if rs.color_hex:
            try:
                run.font.color.rgb = RGBColor(
                    int(rs.color_hex[0:2], 16),
                    int(rs.color_hex[2:4], 16),
                    int(rs.color_hex[4:6], 16),
                )
            except Exception:
                pass
        # Set east-Asian font when text contains CJK characters so they
        # render correctly instead of falling back to a mis-matched glyph.
        if _has_cjk(rs.text):
            self._set_east_asian_font(run, "Microsoft YaHei")

    def _set_east_asian_font(self, run, font_name: str) -> None:
        """Set the east-Asian (CJK) font slot so Chinese characters render correctly."""
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = etree.SubElement(rPr, qn("w:rFonts"))
        rFonts.set(qn("w:eastAsia"), font_name)
        rFonts.set(qn("w:eastAsiaTheme"), "minorEastAsia")

    def _set_bottom_border(self, para, color: str) -> None:
        pPr = para._p.get_or_add_pPr()
        pBdr = pPr.find(qn("w:pBdr"))
        if pBdr is None:
            pBdr = etree.SubElement(pPr, qn("w:pBdr"))
        bottom = pBdr.find(qn("w:bottom"))
        if bottom is None:
            bottom = etree.SubElement(pBdr, qn("w:bottom"))
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), color or "1F3864")

    def _remove(self, elem) -> None:
        parent = elem.getparent()
        if parent is not None:
            parent.remove(elem)
