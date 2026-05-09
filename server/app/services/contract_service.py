import html
import json
import re
import shutil
from collections.abc import AsyncGenerator
from copy import deepcopy
from datetime import datetime
from pathlib import Path
from typing import Any, TypedDict

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from fastapi import HTTPException, status
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import get_settings
from app.models.db_models import CachedLaw, Contract, ContractTemplate
from app.schemas.schemas import ContractGenerateRequest, ContractUpdateRequest, ModelConfig
from app.services.ai_service import AIService
from app.services.boe_service import BoeService

try:
    from langgraph.graph import END, StateGraph
except Exception:  # pragma: no cover
    END = None
    StateGraph = None

try:
    from weasyprint import HTML
except Exception:  # pragma: no cover
    HTML = None

settings = get_settings()

MAX_GENERATION_ATTEMPTS = 3


class WorkflowState(TypedDict):
    db: AsyncSession
    title: str | None
    order_input: str
    model_config: ModelConfig
    extracted_fields: dict[str, str]
    selected_template_id: str | None
    template: ContractTemplate | None
    laws: list[dict[str, str]]
    generated_text: str
    validation_errors: list[str]
    generation_attempts: int


class ContractService:
    def __init__(self) -> None:
        self.boe_service = BoeService()
        self.ai_service = AIService()
        self.graph = self._build_workflow()

    async def generate(self, db: AsyncSession, payload: ContractGenerateRequest) -> Contract:
        initial_state: WorkflowState = {
            "db": db,
            "title": payload.title,
            "order_input": payload.order_input,
            "model_config": payload.llm_config,
            "extracted_fields": {},
            "selected_template_id": None,
            "template": None,
            "laws": [],
            "generated_text": "",
            "validation_errors": [],
            "generation_attempts": 0,
        }
        state = await self._run_workflow(initial_state)
        template = state["template"]
        if template is None:
            raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="Template selection failed")
        if state["validation_errors"]:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Contract validation failed: {', '.join(state['validation_errors'])}",
            )

        contract = Contract(
            title=payload.title or f"Contrato {template.title}",
            template_id=template.id,
            order_input=payload.order_input,
            extracted_fields=state["extracted_fields"],
            generated_text=state["generated_text"],
            laws_used=self._serialize_laws(state["laws"]),
            status="draft",
        )
        db.add(contract)
        await db.commit()
        await db.refresh(contract)
        return contract

    async def generate_stream(
        self, db: AsyncSession, payload: ContractGenerateRequest
    ) -> AsyncGenerator[str, None]:
        def sse(data: dict) -> str:
            return f"data: {json.dumps(data, ensure_ascii=False)}\n\n"

        state: WorkflowState = {
            "db": db,
            "title": payload.title,
            "order_input": payload.order_input,
            "model_config": payload.llm_config,
            "extracted_fields": {},
            "selected_template_id": None,
            "template": None,
            "laws": [],
            "generated_text": "",
            "validation_errors": [],
            "generation_attempts": 0,
        }

        try:
            yield sse({"type": "step", "step": "extract_fields", "status": "running"})
            state = await self._extract_fields_node(state)
            yield sse({"type": "step", "step": "extract_fields", "status": "done",
                       "data": {"fields": state["extracted_fields"]}})

            yield sse({"type": "step", "step": "select_template", "status": "running"})
            state = await self._select_template_node(state)
            template_hint = state["selected_template_id"] or ""
            yield sse({"type": "step", "step": "select_template", "status": "done",
                       "data": {"template_id": template_hint}})

            yield sse({"type": "step", "step": "load_template", "status": "running"})
            state = await self._load_template_node(state)
            template_title = state["template"].title if state["template"] else ""
            yield sse({"type": "step", "step": "load_template", "status": "done",
                       "data": {"template_title": template_title}})

            yield sse({"type": "step", "step": "fetch_laws", "status": "running"})
            state = await self._fetch_laws_node(state)
            yield sse({"type": "step", "step": "fetch_laws", "status": "done",
                       "data": {"laws": [{"boe_id": l["boe_id"], "title": l["title"]} for l in state["laws"]]}})

            attempt = 0
            while True:
                attempt += 1
                yield sse({"type": "step", "step": "generate_contract", "status": "running",
                           "data": {"attempt": attempt}})
                state = await self._generate_contract_node(state)
                yield sse({"type": "step", "step": "generate_contract", "status": "done"})

                yield sse({"type": "step", "step": "validate_contract", "status": "running"})
                state = await self._validate_contract_node(state)
                if state["validation_errors"] and state["generation_attempts"] < MAX_GENERATION_ATTEMPTS:
                    yield sse({"type": "step", "step": "validate_contract", "status": "retrying",
                               "data": {"errors": state["validation_errors"]}})
                    continue
                break

            if state["validation_errors"]:
                yield sse({"type": "step", "step": "validate_contract", "status": "error",
                           "data": {"errors": state["validation_errors"]}})
            else:
                yield sse({"type": "step", "step": "validate_contract", "status": "done"})

            template = state["template"]
            if template is None:
                yield sse({"type": "error", "message": "Template selection failed"})
                return

            contract = Contract(
                title=payload.title or f"Contrato {template.title}",
                template_id=template.id,
                order_input=payload.order_input,
                extracted_fields=state["extracted_fields"],
                generated_text=state["generated_text"],
                laws_used=self._serialize_laws(state["laws"]),
                status="draft",
            )
            db.add(contract)
            await db.commit()
            await db.refresh(contract)

            from app.schemas.schemas import ContractRead
            contract_data = ContractRead.model_validate(contract).model_dump(mode="json")
            yield sse({"type": "done", "contract": contract_data})

        except Exception as e:
            yield sse({"type": "error", "message": str(e)})

    async def list_contracts(self, db: AsyncSession) -> list[Contract]:
        result = await db.scalars(select(Contract).order_by(Contract.created_at.desc()))
        return list(result.all())

    async def get_contract(self, db: AsyncSession, contract_id: str) -> Contract:
        contract = await db.get(Contract, contract_id)
        if not contract:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Contract not found")
        return contract

    async def update_contract(self, db: AsyncSession, contract_id: str, payload: ContractUpdateRequest) -> Contract:
        contract = await self.get_contract(db, contract_id)
        contract.generated_text = payload.generated_text
        contract.status = payload.status
        if payload.title:
            contract.title = payload.title
        await db.commit()
        await db.refresh(contract)
        return contract

    async def delete_contract(self, db: AsyncSession, contract_id: str) -> None:
        contract = await self.get_contract(db, contract_id)
        export_dir = settings.exports_dir / contract.id
        if export_dir.exists():
            shutil.rmtree(export_dir)
        await db.delete(contract)
        await db.commit()

    async def export_docx(self, db: AsyncSession, contract_id: str) -> Path:
        contract = await self.get_contract(db, contract_id)
        file_path = await self._build_docx_file(db, contract, exported=True)

        contract.export_docx_path = str(file_path)
        contract.status = "exported"
        await db.commit()
        return file_path

    async def preview_docx(self, db: AsyncSession, contract_id: str) -> Path:
        contract = await self.get_contract(db, contract_id)
        return await self._build_docx_file(db, contract, exported=False)

    async def _build_docx_file(self, db: AsyncSession, contract: Contract, exported: bool) -> Path:
        export_dir = settings.exports_dir / contract.id
        export_dir.mkdir(parents=True, exist_ok=True)
        file_path = export_dir / ("contract.docx" if exported else "preview.docx")

        template_source = await self._resolve_template_docx_path(db, contract)
        if template_source is not None:
            try:
                self._render_contract_into_template(template_source, file_path, contract)
                return file_path
            except Exception:
                pass

        self._render_plain_docx(file_path, contract)
        return file_path

    def _render_plain_docx(self, file_path: Path, contract: Contract) -> None:
        document = Document()
        document.add_heading(contract.title, level=1)
        for paragraph in (contract.generated_text or "").split("\n\n"):
            clean = paragraph.strip()
            if clean:
                document.add_paragraph(clean)
        document.save(file_path)

    async def export_pdf(self, db: AsyncSession, contract_id: str) -> Path:
        contract = await self.get_contract(db, contract_id)
        if HTML is None:
            raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="WeasyPrint is not available")

        export_dir = settings.exports_dir / contract.id
        export_dir.mkdir(parents=True, exist_ok=True)
        file_path = export_dir / "contract.pdf"
        html_text = html.escape(contract.generated_text or "")
        markup = f"<html><body><h1>{html.escape(contract.title)}</h1><pre>{html_text}</pre></body></html>"
        HTML(string=markup).write_pdf(file_path)

        contract.export_pdf_path = str(file_path)
        contract.status = "exported"
        await db.commit()
        return file_path

    async def _extract_fields(self, order_input: str, model_config: ModelConfig) -> dict[str, str]:
        ai_fields = await self.ai_service.extract_fields(order_input, model_config)
        fields: dict[str, str] = {"raw_order": order_input.strip()}
        if ai_fields:
            fields.update(ai_fields)

        nie_match = re.search(r"NIE[:\s]*([A-Z]\d{7}[A-Z])", order_input, re.IGNORECASE)
        if nie_match:
            fields["nie"] = nie_match.group(1).upper()

        fee_match = re.search(r"(\d+[\.,]?\d*)\s*eu", order_input, re.IGNORECASE)
        if fee_match:
            fields["honorarios"] = fee_match.group(1)

        service_match = re.search(r"服务[:：]\s*([^，,\n]+)|servicio[:\s]*([^，,\n]+)", order_input, re.IGNORECASE)
        if service_match:
            fields["tipo_servicio"] = next(group for group in service_match.groups() if group)

        name_match = re.search(r"客户([^，,\n]+)|cliente[:\s]*([^，,\n]+)", order_input, re.IGNORECASE)
        if name_match:
            fields["nombre"] = next(group for group in name_match.groups() if group).strip()

        return fields

    async def _select_template(
        self, db: AsyncSession, order_input: str, model_config: ModelConfig
    ) -> ContractTemplate:
        result = await db.scalars(select(ContractTemplate).where(ContractTemplate.is_active.is_(True)))
        templates = list(result.all())
        if not templates:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="No active templates available")

        template_options = [
            {"id": t.id, "title": t.title, "category": t.category or "", "subcategory": t.subcategory or ""}
            for t in templates
        ]
        ai_id = await self.ai_service.select_template(order_input, template_options, model_config)
        if ai_id:
            match = next((t for t in templates if t.id == ai_id), None)
            if match:
                return match

        # Fallback: keyword scoring
        order_lower = order_input.lower()
        scored: list[tuple[int, ContractTemplate]] = []
        for template in templates:
            haystack = " ".join(filter(None, [template.title, template.category, template.subcategory or ""])).lower()
            score = sum(1 for token in order_lower.split() if len(token) > 3 and token in haystack)
            scored.append((score, template))
        scored.sort(key=lambda item: item[0], reverse=True)
        return scored[0][1]

    async def _load_template(self, db: AsyncSession, template_id: str) -> ContractTemplate:
        template = await db.get(ContractTemplate, template_id)
        if not template:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Template not found")
        return template

    async def _resolve_laws(self, db: AsyncSession, order_input: str, model_config: ModelConfig) -> list[dict[str, str]]:
        # Extract Latin-script words (5+ chars) as BOE search terms, skip common stopwords
        _STOP = {"para", "como", "esta", "este", "sobre", "entre", "hasta", "desde", "todos", "todas"}
        latin_tokens = [
            w.lower() for w in re.findall(r"[a-zA-ZáéíóúñüÁÉÍÓÚÑÜ]{5,}", order_input)
            if w.lower() not in _STOP
        ]
        # Deduplicate while preserving order
        seen: set[str] = set()
        search_terms: list[str] = []
        for t in latin_tokens:
            if t not in seen:
                seen.add(t)
                search_terms.append(t)
        search_terms = search_terms[:4]

        # Search BOE online for each term and fetch+cache results
        for term in search_terms:
            try:
                boe_results = await self.boe_service.search(term)
                for r in boe_results[:2]:
                    try:
                        await self.boe_service.fetch_and_cache(
                            db,
                            boe_id=r["boe_id"],
                            title=r.get("title"),
                            source_url=r.get("source_url"),
                        )
                    except Exception:
                        pass
            except Exception:
                pass

        # Re-query all cached laws (includes newly fetched ones)
        result = await db.scalars(select(CachedLaw))
        all_laws = list(result.all())

        if not all_laws:
            return []

        # Ask AI to select the most relevant laws
        candidates = [{"boe_id": law.boe_id, "title": law.title} for law in all_laws]
        ai_ids = await self.ai_service.select_laws(order_input, candidates, model_config)
        if ai_ids:
            id_set = set(ai_ids)
            ai_laws = [law for law in all_laws if law.boe_id in id_set]
            if ai_laws:
                return [{"boe_id": law.boe_id, "title": law.title, "raw_text": law.raw_text or ""} for law in ai_laws[:5]]

        # Fallback: keyword scoring
        order_tokens = {w.lower() for w in re.findall(r"[a-zA-ZáéíóúñüÁÉÍÓÚÑÜ]{4,}", order_input)}
        scored: list[tuple[int, CachedLaw]] = []
        for law in all_laws:
            haystack = " ".join(filter(None, [law.title, law.boe_id, law.category or ""])).lower()
            score = sum(1 for token in order_tokens if token in haystack)
            scored.append((score, law))
        scored.sort(key=lambda x: x[0], reverse=True)
        relevant = [law for score, law in scored if score > 0] or [law for _, law in scored[:3]]
        return [{"boe_id": law.boe_id, "title": law.title, "raw_text": law.raw_text or ""} for law in relevant[:5]]

    async def _generate_contract_text(
        self,
        title: str | None,
        order_input: str,
        extracted_fields: dict[str, str],
        template: ContractTemplate,
        laws: list[dict[str, str]],
        model_config: ModelConfig,
    ) -> str:
        ai_text = await self.ai_service.generate_contract(
            order_input=order_input,
            extracted_fields=extracted_fields,
            template_title=template.title,
            template_text=template.raw_text,
            laws=laws,
            model_config=model_config,
        )
        if ai_text:
            return ai_text.strip()

        return self._compose_contract(title, order_input, extracted_fields, template, laws)

    def _compose_contract(
        self,
        title: str | None,
        order_input: str,
        extracted_fields: dict[str, str],
        template: ContractTemplate,
        laws: list[dict[str, str]],
    ) -> str:
        customer = extracted_fields.get("nombre", "la parte cliente")
        service_type = extracted_fields.get("tipo_servicio", template.title)
        fee = extracted_fields.get("honorarios", "por determinar")
        law_text = ", ".join(law["title"] for law in laws)
        template_preview = (template.raw_text or "").splitlines()[:8]
        preview = "\n".join(template_preview)

        return (
            f"{title or template.title}\n\n"
            f"PRIMERA. Partes\n"
            f"Entre el despacho profesional y {customer}, con los datos aportados en el pedido, se formaliza la presente hoja de encargo.\n\n"
            f"SEGUNDA. Objeto\n"
            f"El servicio contratado corresponde a {service_type}. El contenido inicial del pedido es: {order_input.strip()}.\n\n"
            f"TERCERA. Base documental\n"
            f"Se toma como referencia la plantilla '{template.title}' y la normativa siguiente: {law_text}.\n\n"
            f"CUARTA. Honorarios\n"
            f"Los honorarios profesionales pactados ascienden a {fee} EUR, salvo ajuste posterior por acuerdo escrito entre las partes.\n\n"
            f"QUINTA. Clausulas operativas\n"
            f"La parte cliente se compromete a facilitar documentacion veraz y completa. El despacho ejecutara las actuaciones necesarias con diligencia profesional.\n\n"
            f"SEXTA. Referencia de plantilla\n"
            f"Extracto inicial de la plantilla usada:\n{preview}"
        )

    def _validate_contract(self, generated_text: str) -> None:
        required_sections = ["PRIMERA", "CUARTA", "QUINTA"]
        missing = [section for section in required_sections if section not in generated_text]
        if missing:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Generated contract missing sections: {', '.join(missing)}",
            )

    def _validate_contract_sections(self, generated_text: str) -> list[str]:
        required_sections = ["PRIMERA", "CUARTA", "QUINTA"]
        return [section for section in required_sections if section not in generated_text]

    def _serialize_laws(self, laws: list[dict[str, str]]) -> list[dict[str, str]]:
        return [{"boe_id": law["boe_id"], "title": law["title"]} for law in laws]

    async def _resolve_template_docx_path(self, db: AsyncSession, contract: Contract) -> Path | None:
        if not contract.template_id:
            return None

        template = await db.get(ContractTemplate, contract.template_id)
        if template is None or not template.file_path:
            return None

        file_path = Path(template.file_path)
        if not file_path.is_absolute():
            file_path = Path.cwd() / file_path
        if file_path.suffix.lower() != ".docx" or not file_path.exists() or not file_path.is_file():
            return None
        return file_path

    def _render_contract_into_template(self, template_path: Path, output_path: Path, contract: Contract) -> None:
        document = Document(template_path)
        body_start = self._find_template_body_start(document)
        if body_start is None:
            self._render_plain_docx(output_path, contract)
            return

        self._refresh_template_date_line(document, contract, body_start)
        style_samples = self._collect_template_style_samples(document, body_start)
        anchor = document.paragraphs[body_start - 1] if body_start > 0 else None
        body_paragraphs = list(document.paragraphs[body_start:])
        for paragraph in body_paragraphs:
            if self._paragraph_contains_drawing(paragraph):
                anchor = paragraph
                continue
            self._delete_paragraph(paragraph)

        blocks = self._split_contract_blocks(contract.generated_text or "")
        for block in blocks:
            sample = self._select_style_sample(block, style_samples)
            new_paragraph = self._append_paragraph_after(document, anchor)
            self._apply_paragraph_format(sample, new_paragraph)
            self._set_paragraph_text(new_paragraph, block, sample)
            anchor = new_paragraph

        document.save(output_path)

    def _find_template_body_start(self, document: Document) -> int | None:
        markers = ("REUNIDOS", "EXPONEN", "CLÁUSULAS", "CLAUSULAS", "PRIMERA", "SEGUNDA")
        for index, paragraph in enumerate(document.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
            upper = text.upper()
            if any(marker in upper for marker in markers):
                return index

        for index, paragraph in enumerate(document.paragraphs):
            text = paragraph.text.strip()
            if text and len(text) > 40:
                return index
        return None

    def _refresh_template_date_line(self, document: Document, contract: Contract, body_start: int) -> None:
        date_text = self._build_contract_date_text(contract.extracted_fields.get("fecha"))
        for paragraph in document.paragraphs[:body_start]:
            text = paragraph.text.strip()
            if not text:
                continue
            if paragraph.alignment == 2 or text.lower().startswith("en "):
                self._set_paragraph_text(paragraph, date_text, paragraph)
                return

    def _build_contract_date_text(self, raw_date: Any) -> str:
        if isinstance(raw_date, str) and raw_date.strip():
            cleaned = raw_date.strip()
            if cleaned.lower().startswith("en "):
                return cleaned
            return f"En Madrid, a {cleaned}."

        now = datetime.now()
        return f"En Madrid, a {now.day:02d} de {now.month:02d} de {now.year}."

    def _collect_template_style_samples(self, document: Document, body_start: int) -> dict[str, Paragraph]:
        samples: dict[str, Paragraph] = {}
        body_paragraphs = [paragraph for paragraph in document.paragraphs[body_start:] if paragraph.text.strip()]
        if not body_paragraphs:
            fallback = next((paragraph for paragraph in document.paragraphs if paragraph.text.strip()), document.paragraphs[0])
            return {"body": fallback, "section": fallback, "centered": fallback}

        samples["body"] = next((p for p in body_paragraphs if len(p.text.strip()) > 40), body_paragraphs[0])
        samples["section"] = next((p for p in body_paragraphs if self._looks_like_section_heading(p.text.strip())), samples["body"])
        samples["centered"] = next(
            (p for p in body_paragraphs if p.alignment == 1 and len(p.text.strip()) <= 40),
            samples["section"],
        )
        return samples

    def _split_contract_blocks(self, text: str) -> list[str]:
        normalized = text.replace("\r\n", "\n").strip()
        if not normalized:
            return []

        blocks: list[str] = []
        for chunk in re.split(r"\n\s*\n+", normalized):
            lines = [line.strip() for line in chunk.split("\n") if line.strip()]
            blocks.extend(lines)
        return blocks

    def _select_style_sample(self, block: str, samples: dict[str, Paragraph]) -> Paragraph:
        if self._looks_like_centered_heading(block):
            return samples["centered"]
        if self._looks_like_section_heading(block):
            return samples["section"]
        return samples["body"]

    def _looks_like_centered_heading(self, text: str) -> bool:
        normalized = text.strip()
        return bool(normalized) and len(normalized) <= 40 and normalized == normalized.upper()

    def _looks_like_section_heading(self, text: str) -> bool:
        normalized = text.strip().upper()
        return bool(
            re.match(
                r"^(PRIMERA|SEGUNDA|TERCERA|CUARTA|QUINTA|SEXTA|SEPTIMA|S[EÉ]PTIMA|OCTAVA|NOVENA|D[EÉ]CIMA|UND[EÉ]CIMA|DUOD[EÉ]CIMA|DECIMOTERCERA)",
                normalized,
            )
        )

    def _append_paragraph_after(self, document: Document, anchor: Paragraph | None) -> Paragraph:
        if anchor is None:
            return document.add_paragraph()
        paragraph = document.add_paragraph()
        anchor._p.addnext(paragraph._p)
        return paragraph

    def _apply_paragraph_format(self, source: Paragraph, target: Paragraph) -> None:
        if source._p.pPr is not None:
            target._p.insert(0, deepcopy(source._p.pPr))

    def _set_paragraph_text(self, paragraph: Paragraph, text: str, sample: Paragraph) -> None:
        for child in list(paragraph._p):
            if child.tag != qn("w:pPr"):
                paragraph._p.remove(child)
        run = paragraph.add_run(text)
        if sample.runs and sample.runs[0]._element.rPr is not None:
            run._element.insert(0, deepcopy(sample.runs[0]._element.rPr))

    def _paragraph_contains_drawing(self, paragraph: Paragraph) -> bool:
        return bool(paragraph._p.xpath(".//*[local-name()='drawing' or local-name()='pict']"))

    def _delete_paragraph(self, paragraph: Paragraph) -> None:
        element = paragraph._element
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)

    async def _run_workflow(self, state: WorkflowState) -> WorkflowState:
        if self.graph is not None:
            return await self.graph.ainvoke(state)
        return await self._run_workflow_fallback(state)

    def _build_workflow(self):
        if StateGraph is None or END is None:
            return None

        workflow = StateGraph(WorkflowState)
        workflow.add_node("extract_fields", self._extract_fields_node)
        workflow.add_node("select_template", self._select_template_node)
        workflow.add_node("load_template", self._load_template_node)
        workflow.add_node("fetch_laws", self._fetch_laws_node)
        workflow.add_node("generate_contract", self._generate_contract_node)
        workflow.add_node("validate_contract", self._validate_contract_node)
        workflow.set_entry_point("extract_fields")
        workflow.add_edge("extract_fields", "select_template")
        workflow.add_edge("select_template", "load_template")
        workflow.add_edge("load_template", "fetch_laws")
        workflow.add_edge("fetch_laws", "generate_contract")
        workflow.add_edge("generate_contract", "validate_contract")
        workflow.add_conditional_edges(
            "validate_contract",
            self._workflow_after_validate,
            {
                "retry_generate": "generate_contract",
                "end": END,
            },
        )
        return workflow.compile()

    async def _run_workflow_fallback(self, state: WorkflowState) -> WorkflowState:
        current_state = await self._extract_fields_node(state)
        current_state = await self._select_template_node(current_state)
        current_state = await self._load_template_node(current_state)
        current_state = await self._fetch_laws_node(current_state)
        while True:
            current_state = await self._generate_contract_node(current_state)
            current_state = await self._validate_contract_node(current_state)
            if self._workflow_after_validate(current_state) == "end":
                return current_state

    async def _extract_fields_node(self, state: WorkflowState) -> WorkflowState:
        extracted_fields = await self._extract_fields(state["order_input"], state["model_config"])
        state["extracted_fields"] = extracted_fields
        return state

    async def _select_template_node(self, state: WorkflowState) -> WorkflowState:
        template = await self._select_template(state["db"], state["order_input"], state["model_config"])
        state["selected_template_id"] = template.id
        return state

    async def _load_template_node(self, state: WorkflowState) -> WorkflowState:
        template_id = state["selected_template_id"]
        if not template_id:
            raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="Template id missing")
        state["template"] = await self._load_template(state["db"], template_id)
        return state

    async def _fetch_laws_node(self, state: WorkflowState) -> WorkflowState:
        state["laws"] = await self._resolve_laws(state["db"], state["order_input"], state["model_config"])
        return state

    async def _generate_contract_node(self, state: WorkflowState) -> WorkflowState:
        template = state["template"]
        if template is None:
            raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="Template not loaded")
        state["generation_attempts"] += 1
        state["generated_text"] = await self._generate_contract_text(
            state["title"],
            state["order_input"],
            state["extracted_fields"],
            template,
            state["laws"],
            state["model_config"],
        )
        return state

    async def _validate_contract_node(self, state: WorkflowState) -> WorkflowState:
        state["validation_errors"] = self._validate_contract_sections(state["generated_text"])
        return state

    def _workflow_after_validate(self, state: WorkflowState) -> str:
        if state["validation_errors"] and state["generation_attempts"] < MAX_GENERATION_ATTEMPTS:
            return "retry_generate"
        return "end"
