from pathlib import Path

import pdfplumber
from docx import Document


class DocumentParser:
    @staticmethod
    def parse(file_path: Path) -> str:
        suffix = file_path.suffix.lower()
        if suffix == ".docx":
            return DocumentParser._parse_docx(file_path)
        if suffix == ".pdf":
            return DocumentParser._parse_pdf(file_path)
        raise ValueError(f"Unsupported file type: {suffix}")

    @staticmethod
    def _parse_docx(file_path: Path) -> str:
        document = Document(file_path)
        parts: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    @staticmethod
    def _parse_pdf(file_path: Path) -> str:
        pages: list[str] = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = (page.extract_text() or "").strip()
                if text:
                    pages.append(text)
        return "\n\n".join(pages)
