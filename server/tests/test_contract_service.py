import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.models.db_models import ContractTemplate
from app.schemas.schemas import ContractGenerateRequest
from app.services.contract_service import ContractService


@pytest.mark.asyncio
async def test_generate_contract_falls_back_and_persists(db_session, model_config):
    template = ContractTemplate(
        title="HE Reagrupacion familiar",
        category="EXTRANJERIA",
        subcategory="REAGRUPACION FAMILIAR",
        file_name="reagrupacion.docx",
        file_path="/tmp/reagrupacion.docx",
        raw_text="Plantilla base\nPRIMERA\nSEGUNDA\nCUARTA\nQUINTA",
        language="es",
        is_active=True,
    )
    db_session.add(template)
    await db_session.commit()

    service = ContractService()

    async def fake_ai_generate_contract(*args, **kwargs):
        return None

    service.ai_service.generate_contract = fake_ai_generate_contract

    payload = ContractGenerateRequest(
        title="Contrato de prueba",
        order_input="客户王芳，NIE: Y9876543B，服务：家庭团聚居留，费用：1900欧。",
        model_config=model_config,
    )

    contract = await service.generate(db_session, payload)

    assert contract.id
    assert contract.template_id == template.id
    assert contract.title == "Contrato de prueba"
    assert "PRIMERA. Partes" in (contract.generated_text or "")
    assert "CUARTA. Honorarios" in (contract.generated_text or "")
    assert contract.laws_used


@pytest.mark.asyncio
async def test_export_docx_contains_generated_contract_text(db_session, model_config, tmp_path):
    template_path = tmp_path / "template.docx"
    document = Document()
    title_paragraph = document.add_paragraph("HOJA DE ENCARGO PROFESIONAL")
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_paragraph = document.add_paragraph("En Madrid, a 23 de 09 de 2025.")
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    reunidos = document.add_paragraph("REUNIDOS")
    reunidos.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro = document.add_paragraph("Texto antiguo de introduccion.")
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    clausulas = document.add_paragraph("CLÁUSULAS")
    clausulas.alignment = WD_ALIGN_PARAGRAPH.CENTER
    primera = document.add_paragraph("PRIMERA.- Texto antiguo de plantilla.")
    primera.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    segunda = document.add_paragraph("SEGUNDA.- Texto antiguo de plantilla.")
    segunda.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    document.save(template_path)

    template = ContractTemplate(
        title="HE Formato original",
        category="MERCANTIL",
        subcategory="GENERAL",
        file_name="template.docx",
        file_path=str(template_path),
        raw_text="Plantilla Original\nTexto con formato de plantilla.",
        language="es",
        is_active=True,
    )
    db_session.add(template)
    await db_session.commit()

    service = ContractService()

    async def fake_ai_generate_contract(*args, **kwargs):
        return (
            "REUNIDOS\n"
            "Nueva introduccion del contrato.\n\n"
            "CLÁUSULAS\n"
            "PRIMERA.- Texto nuevo de prueba.\n"
            "SEGUNDA.- Honorarios pactados en 1900 EUR.\n"
            "TERCERA.- Clausulas operativas de prueba.\n"
            "CUARTA.- Referencias documentales de prueba.\n"
            "QUINTA.- Cierre contractual de prueba."
        )

    service.ai_service.generate_contract = fake_ai_generate_contract

    payload = ContractGenerateRequest(
        title="Contrato formato",
        order_input="客户王芳，NIE: Y9876543B，服务：家庭团聚居留，费用：1900欧。",
        model_config=model_config,
    )

    contract = await service.generate(db_session, payload)
    preview_path = await service.preview_docx(db_session, contract.id)
    export_path = await service.export_docx(db_session, contract.id)
    preview_doc = Document(preview_path)
    exported_doc = Document(export_path)
    exported_paragraphs = [paragraph for paragraph in exported_doc.paragraphs if paragraph.text.strip()]
    preview_paragraphs = [paragraph for paragraph in preview_doc.paragraphs if paragraph.text.strip()]
    exported_text = "\n".join(paragraph.text for paragraph in exported_paragraphs)

    assert exported_paragraphs[0].text == "HOJA DE ENCARGO PROFESIONAL"
    assert exported_paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert exported_paragraphs[1].text.startswith("En Madrid, a ")
    assert exported_paragraphs[1].alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert "Nueva introduccion del contrato." in exported_text
    assert "SEGUNDA.- Honorarios pactados en 1900 EUR." in exported_text
    assert "Texto antiguo de plantilla." not in exported_text
    assert preview_paragraphs[0].text == "HOJA DE ENCARGO PROFESIONAL"
